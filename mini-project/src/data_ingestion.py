#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Data Ingestion Module for Resume-Job Matching System.

This module provides classes for reading resume files (PDF/DOCX) and job postings (CSV).
"""

import os
import logging
import pandas as pd
from pathlib import Path
from typing import Dict, List, Any, Union

# For PDF parsing
import fitz  # PyMuPDF

# For DOCX parsing
import docx

logger = logging.getLogger(__name__)


class ResumeReader:
    """Class for reading resume files from a directory."""
    
    def __init__(self, resume_dir: Union[str, Path]):
        """Initialize with the directory containing resume files.
        
        Args:
            resume_dir: Directory path containing resume files (PDF/DOCX)
        """
        self.resume_dir = Path(resume_dir)
        logger.info(f"Initialized ResumeReader with directory: {self.resume_dir}")
    
    def read_resume(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Read a single resume file.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing resume data with keys:
                - file_path: Path to the resume file
                - file_name: Name of the resume file
                - content: Raw text content of the resume
                - file_type: Type of the file (pdf/docx)
        """
        file_path = Path(file_path)
        try:
            file_name = file_path.name
            file_type = file_path.suffix.lower()[1:]  # Remove the dot
            
            # Read the content based on file type
            if file_type == "pdf":
                content = self._read_pdf(file_path)
            elif file_type == "docx":
                content = self._read_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            resume_data = {
                "file_path": str(file_path),
                "file_name": file_name,
                "content": content,
                "file_type": file_type
            }
            
            logger.info(f"Successfully read resume: {file_name}")
            return resume_data
            
        except Exception as e:
            logger.error(f"Error reading resume {file_path}: {str(e)}")
            raise
    
    def read_resumes(self) -> List[Dict[str, Any]]:
        """Read all resume files from the directory.
        
        Returns:
            List of dictionaries containing resume data with keys:
                - file_path: Path to the resume file
                - file_name: Name of the resume file
                - content: Raw text content of the resume
                - file_type: Type of the file (pdf/docx)
        """
        resumes = []
        
        if not self.resume_dir.exists():
            logger.warning(f"Resume directory {self.resume_dir} does not exist")
            return resumes
        
        # Get all PDF and DOCX files
        resume_files = list(self.resume_dir.glob("*.pdf")) + list(self.resume_dir.glob("*.docx"))
        
        logger.info(f"Found {len(resume_files)} resume files")
        
        for file_path in resume_files:
            try:
                file_name = file_path.name
                file_type = file_path.suffix.lower()[1:]  # Remove the dot
                
                # Read the content based on file type
                if file_type == "pdf":
                    content = self._read_pdf(file_path)
                elif file_type == "docx":
                    content = self._read_docx(file_path)
                else:
                    logger.warning(f"Unsupported file type: {file_type} for {file_path}")
                    continue
                
                resumes.append({
                    "file_path": str(file_path),
                    "file_name": file_name,
                    "content": content,
                    "file_type": file_type
                })
                
                logger.debug(f"Successfully read resume: {file_name}")
                
            except Exception as e:
                logger.error(f"Error reading resume {file_path}: {str(e)}")
        
        logger.info(f"Successfully read {len(resumes)} resumes")
        return resumes
    
    def _read_pdf(self, file_path: Path) -> str:
        """Read text content from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Text content of the PDF file
        """
        try:
            text = ""
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text()
            return text
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            raise
    
    def _read_docx(self, file_path: Path) -> str:
        """Read text content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Text content of the DOCX file
        """
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            raise


class JobReader:
    """Class for reading job postings from a CSV file."""
    
    def __init__(self, jobs_file: Union[str, Path] = None):
        """Initialize with the path to the CSV file containing job postings.
        
        Args:
            jobs_file: Optional path to the CSV file containing job postings
        """
        self.jobs_file = Path(jobs_file) if jobs_file else None
        if self.jobs_file:
            logger.info(f"Initialized JobReader with file: {self.jobs_file}")
        else:
            logger.info("Initialized JobReader without default jobs file")
    
    def read_jobs(self, jobs_file: Union[str, Path] = None) -> List[Dict[str, Any]]:
        """Read job postings from the CSV file.
        
        Args:
            jobs_file: Optional path to the CSV file. If not provided, uses the file
                      specified during initialization.
        
        Returns:
            List of dictionaries containing job data
        """
        jobs = []
        file_path = Path(jobs_file) if jobs_file else self.jobs_file
        
        if not file_path:
            logger.warning("No jobs file specified")
            return jobs
        
        if not file_path.exists():
            logger.warning(f"Jobs file {file_path} does not exist")
            return jobs
        
        try:
            # Read the CSV file
            df = pd.read_csv(file_path)
            
            # Check if required columns exist
            required_columns = ["title", "location", "required_skills", "job_description", "salary"]
            missing_columns = [col for col in required_columns if col not in df.columns]
            
            if missing_columns:
                logger.warning(f"Missing columns in jobs file: {missing_columns}")
                # Add missing columns with empty values
                for col in missing_columns:
                    df[col] = ""
            
            # Convert DataFrame to list of dictionaries
            jobs = df.to_dict(orient="records")
            
            # Process required_skills to convert from string to list if needed
            for job in jobs:
                if "required_skills" in job and isinstance(job["required_skills"], str):
                    job["required_skills"] = [skill.strip() for skill in job["required_skills"].split(",")]
                elif "required_skills" not in job:
                    job["required_skills"] = []
            
            logger.info(f"Successfully read {len(jobs)} jobs from {file_path}")
            
        except Exception as e:
            logger.error(f"Error reading jobs file {file_path}: {str(e)}")
        
        return jobs


class DataIngestion:
    """Main class for data ingestion that wraps ResumeReader and JobReader functionality."""
    
    def __init__(self):
        """Initialize the DataIngestion class."""
        self.resume_reader = None
        self.job_reader = None
        logger.info("Initialized DataIngestion")
    
    def load_resume(self, resume_path: Union[str, Path]) -> str:
        """Load a single resume file.
        
        Args:
            resume_path: Path to the resume file
            
        Returns:
            Text content of the resume
        """
        resume_path = Path(resume_path)
        if not resume_path.exists():
            raise FileNotFoundError(f"Resume file not found: {resume_path}")
            
        try:
            if resume_path.suffix.lower() == '.pdf':
                return self._read_pdf(resume_path)
            elif resume_path.suffix.lower() == '.docx':
                return self._read_docx(resume_path)
            else:
                raise ValueError(f"Unsupported file type: {resume_path.suffix}")
        except Exception as e:
            logger.error(f"Error loading resume {resume_path}: {str(e)}")
            raise
    
    def load_jobs(self, jobs_path: Union[str, Path]) -> List[Dict[str, Any]]:
        """Load job postings from a CSV file.
        
        Args:
            jobs_path: Path to the jobs CSV file
            
        Returns:
            List of dictionaries containing job data
        """
        self.job_reader = JobReader(jobs_path)
        return self.job_reader.read_jobs()
    
    def _read_pdf(self, file_path: Path) -> str:
        """Read text content from a PDF file."""
        try:
            text = ""
            with fitz.open(file_path) as pdf:
                for page in pdf:
                    text += page.get_text()
            return text
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            raise
    
    def _read_docx(self, file_path: Path) -> str:
        """Read text content from a DOCX file."""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            raise


# Create sample data for testing
def create_sample_data(base_dir: Union[str, Path]):
    """Create sample data for testing.
    
    Args:
        base_dir: Base directory for the project
    """
    base_dir = Path(base_dir)
    
    # Create sample jobs CSV
    jobs_file = base_dir / "data" / "dummy_data" / "sample_jobs.csv"
    
    sample_jobs = pd.DataFrame({
        "title": [
            "Data Scientist", 
            "Machine Learning Engineer", 
            "Software Developer", 
            "Frontend Developer", 
            "Backend Developer"
        ],
        "location": [
            "New York, NY", 
            "San Francisco, CA", 
            "Austin, TX", 
            "Seattle, WA", 
            "Boston, MA"
        ],
        "required_skills": [
            "Python, SQL, Machine Learning, Statistics",
            "Python, TensorFlow, PyTorch, Computer Vision",
            "Java, Spring, Hibernate, SQL",
            "JavaScript, React, HTML, CSS",
            "Python, Django, Flask, PostgreSQL"
        ],
        "job_description": [
            "We are looking for a Data Scientist to join our team. You will be responsible for analyzing data and building models to solve business problems.",
            "We are looking for a Machine Learning Engineer to join our team. You will be responsible for building and deploying machine learning models.",
            "We are looking for a Software Developer to join our team. You will be responsible for building and maintaining our core applications.",
            "We are looking for a Frontend Developer to join our team. You will be responsible for building and maintaining our user interfaces.",
            "We are looking for a Backend Developer to join our team. You will be responsible for building and maintaining our server-side applications."
        ],
        "salary": [
            "$120,000 - $150,000",
            "$130,000 - $160,000",
            "$100,000 - $130,000",
            "$90,000 - $120,000",
            "$110,000 - $140,000"
        ]
    })
    
    # Create directory if it doesn't exist
    os.makedirs(jobs_file.parent, exist_ok=True)
    
    # Save to CSV
    sample_jobs.to_csv(jobs_file, index=False)
    logger.info(f"Created sample jobs file: {jobs_file}")


if __name__ == "__main__":
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Create sample data
    create_sample_data(Path(__file__).parent.parent)
    
    # Test reading jobs
    job_reader = JobReader(Path(__file__).parent.parent / "data" / "dummy_data" / "sample_jobs.csv")
    jobs = job_reader.read_jobs()
    
    print(f"Read {len(jobs)} jobs")
    for i, job in enumerate(jobs[:2]):  # Print first 2 jobs
        print(f"\nJob {i+1}:")
        print(f"Title: {job['title']}")
        print(f"Location: {job['location']}")
        print(f"Required Skills: {job['required_skills']}")
        print(f"Salary: {job['salary']}")